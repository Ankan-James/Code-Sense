"""
CodeSense v2 — Tools
Code parser (Python AST + AI for other languages) and report generators (PDF + DOCX).
"""

import ast
import os
import json
import tempfile


# ─── Language Detection ────────────────────────────────────────────────────────

EXTENSION_MAP = {
    ".py": "Python",
    ".js": "JavaScript",
    ".ts": "TypeScript",
    ".java": "Java",
    ".cs": "C#",
    ".c": "C",
    ".cpp": "C++",
    ".html": "HTML",
    ".css": "CSS",
}


def detect_language(filename):
    """Detects the programming language from file extension."""
    ext = os.path.splitext(filename)[1].lower()
    return EXTENSION_MAP.get(ext, "Unknown")


# ─── Master Segmentation Prompt ───────────────────────────────────────────────

SEGMENT_SYSTEM_PROMPT = """You are a code parser for the CodeSense application.
Your only job is to receive a code file and return a recursive JSON segment tree.
You must follow the exact language rules and JSON shape you have been given.
You must never return markdown, backticks, explanations, or any text outside the JSON array.
Your entire response must be a valid JSON array that can be parsed by JSON.parse() with no modifications.
If you are unsure about a segment, make your best judgment based on the rules.
Never skip segments. Never merge segments that should be separate.
Always include the actual source code in the code field, not a description.

The JSON shape (always return this exact structure):
[
  {
    "name": "segment name",
    "type": "segment type",
    "code": "the actual code of this segment",
    "children": [
      {
        "name": "child segment name",
        "type": "child segment type",
        "code": "the actual code of this child segment",
        "children": []
      }
    ]
  }
]

Rules for the JSON shape:
- Every segment must have exactly four keys: name, type, code, children
- children is always a list — empty list [] if the segment has no children
- children can contain more segments which can also have their own children — this is recursive and can go as deep as the code requires
- code must contain the actual source code of that segment, not a description
- name must be a short meaningful identifier — the function name, class name, or block label
- type must be one of the values defined in the language rules
- Never return markdown, never return backticks, never return explanations — return only the raw JSON array

Language Rules:

### JavaScript
Top level segments in this order:
1. type: "imports" — all import and require statements grouped together, name: "Imports"
2. type: "constants" — all top level const, let, var declarations that are not functions, name: "Constants"
3. type: "class" — each class is its own segment, name: the class name
   - children: each method inside the class is a child segment, type: "method", name: the method name, children: []
4. type: "function" — each top level function or arrow function assigned to a variable, name: the function name
   - children: any nested functions inside, type: "function", name: the nested function name, children: []
5. type: "event_listener" — all addEventListener or DOMContentLoaded blocks, name: "Event Listeners"
6. type: "init" — the main init or entry call at the bottom of the file, name: "Init"

### TypeScript
Top level segments in this order:
1. type: "imports" — all import statements, name: "Imports"
2. type: "types" — all interface and type definitions grouped, name: "Types and Interfaces"
   - children: each individual interface or type is a child, type: "interface" or "type", name: the interface or type name, children: []
3. type: "enums" — all enum definitions, name: "Enums"
   - children: each enum is a child, type: "enum", name: the enum name, children: []
4. type: "constants" — all top level const declarations that are not functions, name: "Constants"
5. type: "class" — each class is its own segment, name: the class name
   - children: each method inside the class, type: "method", name: the method name, children: []
6. type: "function" — each top level function, name: the function name
   - children: any nested functions, type: "function", name: the nested function name, children: []
7. type: "event_listener" — all event listener blocks, name: "Event Listeners"
8. type: "init" — the main init or entry call, name: "Init"

### Java
Top level segments in this order:
1. type: "package" — the package declaration, name: "Package"
2. type: "imports" — all import statements grouped, name: "Imports"
3. type: "class" — the main class, name: the class name
   - children in this order:
     - type: "constants" — all static final fields, name: "Constants", children: []
     - type: "fields" — all instance fields, name: "Fields", children: []
     - type: "constructor" — the constructor method, name: the constructor name, children: []
     - type: "method" — each public method, name: the method name, children: []
     - type: "method" — each private method, name: the method name, children: []
     - type: "method" — each static method, name: the method name, children: []
     - type: "main" — the main() method if present, name: "main", children: []

### C#
Top level segments in this order:
1. type: "using" — all using directives grouped, name: "Using Directives"
2. type: "namespace" — the namespace block, name: the namespace name
   - children in this order:
     - type: "enum" — each enum definition, name: the enum name, children: []
     - type: "interface" — each interface definition, name: the interface name, children: []
     - type: "class" — each class definition, name: the class name
       - children in this order:
         - type: "constants" — all const fields, name: "Constants", children: []
         - type: "fields" — all instance fields, name: "Fields", children: []
         - type: "property" — each property with get/set, name: the property name
           - children: type: "getter" name: "get" children: [] and type: "setter" name: "set" children: [] if both exist
         - type: "constructor" — the constructor, name: the constructor name, children: []
         - type: "method" — each public method, name: the method name, children: []
         - type: "method" — each private method, name: the method name, children: []
         - type: "method" — each static method, name: the method name, children: []
         - type: "main" — the Main() entry point if present, name: "Main", children: []

### C
Top level segments in this order:
1. type: "preprocessor" — all #include and #define statements grouped, name: "Preprocessor"
   - children: each #include is a child type: "include" name: the header name children: []
   - children: each #define is a child type: "define" name: the macro name children: []
2. type: "typedef" — all struct and typedef definitions, name: "Type Definitions"
   - children: each struct or typedef is a child, type: "struct" or "typedef", name: the type name, children: []
3. type: "globals" — all global variable declarations, name: "Global Variables", children: []
4. type: "prototype" — all function prototypes grouped, name: "Function Prototypes", children: []
5. type: "function" — each function definition, name: the function name, children: []
6. type: "main" — the main() function, name: "main", children: []

### C++
Top level segments in this order:
1. type: "preprocessor" — all #include and using namespace statements, name: "Preprocessor"
   - children: each #include is a child type: "include" name: the header name children: []
   - children: using namespace is a child type: "namespace" name: the namespace name children: []
2. type: "constants" — all top level const and constexpr declarations, name: "Constants", children: []
3. type: "typedef" — all struct and typedef definitions, name: "Type Definitions"
   - children: each struct or typedef, type: "struct" or "typedef", name: the type name, children: []
4. type: "globals" — all global variable declarations, name: "Global Variables", children: []
5. type: "class" — each class definition, name: the class name
   - children in this order:
     - type: "private" — private section, name: "Private", children: []
     - type: "public" — public section, name: "Public"
       - children: constructor type: "constructor" name: the class name children: []
       - children: destructor type: "destructor" name: ~ClassName children: []
       - children: each method type: "method" name: the method name children: []
     - type: "method_def" — out of class method definitions ClassName::method(), name: the full method signature, children: []
6. type: "prototype" — all function prototypes, name: "Function Prototypes", children: []
7. type: "function" — each standalone function, name: the function name, children: []
8. type: "main" — the main() function, name: "main", children: []

### HTML
Top level segments in this order:
1. type: "doctype" — the DOCTYPE declaration, name: "DOCTYPE", children: []
2. type: "head" — the entire head element, name: "Head"
   - children: each meaningful child of head
     - type: "meta" — all meta tags grouped, name: "Meta Tags", children: []
     - type: "title" — the title tag, name: "Title", children: []
     - type: "stylesheet" — each link rel stylesheet, name: the href value, children: []
     - type: "script" — each script tag in head, name: the src value or "Inline Script", children: []
3. type: "body" — the entire body element, name: "Body"
   - children: each meaningful semantic child of body
     - type: "header" name: "Header"
       - children: type: "nav" name: "Nav" children: []
     - type: "main" name: "Main"
       - children: each section or article or meaningful div
         - type: "section" name: the id or class of the section
         - type: "article" name: the id or class of the article
     - type: "footer" name: "Footer", children: []
     - type: "script" — each script tag at bottom of body, name: the src value or "Inline Script", children: []

Rule for HTML: Stop drilling down at meaningful semantic tags. Do not go down to every span, a, p, or button.

### CSS
Top level segments — each rule block is a segment:
1. type: "reset" — the * {} rule if present, name: "Reset", children: []
2. type: "variables" — the :root {} block, name: "Variables", children: []
3. type: "base" — body, html base rules, name: "Base Styles", children: []
4. type: "component" — each class or id selector block, name: the selector name, children: []
5. type: "keyframes" — each @keyframes block, name: the animation name
   - children: type: "keyframe" name: "from" or "to" or the percentage, children: []
6. type: "media" — each @media block, name: the media query condition
   - children: each selector inside the media query, type: "component", name: the selector name, children: []

Rule for CSS: A segment has children if and only if its block contains other {} blocks inside it. Otherwise children is always [].
"""


# ─── Segment Validation ───────────────────────────────────────────────────────

def validate_segments(segments):
    """Recursively validates and cleans parsed segments to ensure the correct shape."""
    validated = []
    for seg in segments:
        if not isinstance(seg, dict):
            continue
        clean = {
            "name": seg.get("name", "Unknown"),
            "type": seg.get("type", "block"),
            "code": seg.get("code", ""),
            "children": [],
        }
        # Recursively validate children
        raw_children = seg.get("children", [])
        if isinstance(raw_children, list):
            clean["children"] = validate_segments(raw_children)
        validated.append(clean)
    return validated


# ─── Code Parsing ──────────────────────────────────────────────────────────────

def get_child_functions(node, code):
    """Extracts nested/child functions and methods from a parent node."""
    children = []
    for child in ast.iter_child_nodes(node):
        if isinstance(child, ast.FunctionDef) or isinstance(child, ast.AsyncFunctionDef):
            child_source = ast.get_source_segment(code, child)
            children.append({
                "type": "method" if isinstance(node, ast.ClassDef) else "function",
                "name": child.name,
                "code": child_source or "",
                "children": [],
            })
    return children


def parse_python(code):
    """Parses Python code using AST into the standard segment tree shape."""
    segments = []

    try:
        tree = ast.parse(code)
    except SyntaxError:
        return [{
            "type": "full_code",
            "name": "Full File",
            "code": code,
            "children": [],
        }]

    # --- Collect imports ---
    import_lines = []
    for node in tree.body:
        if isinstance(node, (ast.Import, ast.ImportFrom)):
            source = ast.get_source_segment(code, node)
            if source:
                import_lines.append(source)

    if import_lines:
        segments.append({
            "type": "imports",
            "name": "Imports",
            "code": "\n".join(import_lines),
            "children": [],
        })

    # --- Collect top-level constants/assignments ---
    constant_lines = []
    for node in tree.body:
        if isinstance(node, ast.Assign):
            source = ast.get_source_segment(code, node)
            if source:
                constant_lines.append(source)
        elif isinstance(node, ast.AnnAssign):
            source = ast.get_source_segment(code, node)
            if source:
                constant_lines.append(source)

    if constant_lines:
        segments.append({
            "type": "constants",
            "name": "Constants",
            "code": "\n".join(constant_lines),
            "children": [],
        })

    # --- Collect classes and functions ---
    for node in tree.body:
        if isinstance(node, ast.ClassDef):
            segments.append({
                "type": "class",
                "name": node.name,
                "code": ast.get_source_segment(code, node) or "",
                "children": get_child_functions(node, code),
            })
        elif isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)):
            segments.append({
                "type": "function",
                "name": node.name,
                "code": ast.get_source_segment(code, node) or "",
                "children": get_child_functions(node, code),
            })

    # Fallback if nothing was extracted
    if not segments:
        segments.append({
            "type": "full_code",
            "name": "Full File",
            "code": code,
            "children": [],
        })

    return segments


def parse_non_python(code, filename, client):
    """Uses AI to parse non-Python code into a recursive segment tree."""
    language = detect_language(filename)

    user_prompt = f"""Parse this {language} file into a recursive segment tree.
Follow the {language} rules exactly.
Return only a raw JSON array. No markdown. No backticks. No explanation. No text before or after the JSON.

File content:
{code}"""

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system", "content": SEGMENT_SYSTEM_PROMPT},
            {"role": "user", "content": user_prompt},
        ]
    )

    try:
        raw = response.choices[0].message.content.strip()
        # Strip markdown fences if AI accidentally wraps it
        if raw.startswith("```"):
            raw = raw.split("\n", 1)[1] if "\n" in raw else raw[3:]
            if raw.endswith("```"):
                raw = raw[:-3]
            raw = raw.strip()

        result = json.loads(raw)
        return validate_segments(result)
    except Exception:
        return [{
            "type": "full_code",
            "name": "Full File",
            "code": code,
            "children": [],
        }]


def parse_code(code, filename="", client=None):
    """Breaks code into a nested tree of segments."""
    if filename.endswith(".py"):
        return parse_python(code)
    else:
        if client:
            return parse_non_python(code, filename, client)
        else:
            return [{
                "type": "full_code",
                "name": "Full File",
                "code": code,
                "children": [],
            }]


# ─── Report Generation ─────────────────────────────────────────────────────────

def generate_pdf_report(segments, personality, filename):
    """Generates a PDF report of all segment explanations."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 20)
    pdf.cell(0, 15, "CodeSense Report", ln=True, align="C")
    pdf.set_font("Helvetica", "", 12)
    pdf.cell(0, 10, f"File: {filename}", ln=True, align="C")
    pdf.cell(0, 8, f"Tutor Personality: {personality}", ln=True, align="C")
    pdf.ln(10)

    # Segments
    def write_segments(segs, depth=0):
        indent = "  " * depth
        for segment in segs:
            # Segment header
            pdf.set_font("Helvetica", "B", max(14 - depth * 2, 10))
            segment_title = f"{indent}{segment.get('type', 'block').upper()} - {segment.get('name', 'Unknown')}"
            pdf.cell(0, 10, segment_title, ln=True)
            pdf.ln(2)

            # Segment code
            pdf.set_font("Courier", "", 9)
            code = segment.get("code", "")
            for line in code.split("\n"):
                safe_line = line.encode("latin-1", "replace").decode("latin-1")
                pdf.cell(0, 5, f"  {indent}{safe_line}", ln=True)
            pdf.ln(4)

            # Explanation
            explanation = segment.get("explanation", "No explanation generated.")
            pdf.set_font("Helvetica", "", 11)
            safe_explanation = explanation.encode("latin-1", "replace").decode("latin-1")
            pdf.multi_cell(0, 6, safe_explanation)
            pdf.ln(8)

            # Recurse into children
            children = segment.get("children", [])
            if children:
                write_segments(children, depth + 1)

    write_segments(segments)

    # Save to temp file
    report_path = os.path.join(tempfile.gettempdir(), f"codesense_report_{filename}.pdf")
    pdf.output(report_path)
    return report_path


def generate_docx_report(segments, personality, filename):
    """Generates a Word document report of all segment explanations."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading("CodeSense Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"File: {filename}\nTutor Personality: {personality}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(136, 136, 136)

    doc.add_paragraph("")  # spacer

    # Segments
    def write_segments(segs, depth=0):
        for segment in segs:
            # Segment header
            heading_level = min(2 + depth, 4)
            doc.add_heading(
                f"{segment.get('type', 'block').upper()} — {segment.get('name', 'Unknown')}",
                level=heading_level
            )

            # Code block
            code_para = doc.add_paragraph()
            code_run = code_para.add_run(segment.get("code", ""))
            code_run.font.name = "Courier New"
            code_run.font.size = Pt(9)
            code_run.font.color.rgb = RGBColor(200, 200, 200)

            # Explanation
            explanation = segment.get("explanation", "No explanation generated.")
            doc.add_paragraph(explanation)
            doc.add_paragraph("")  # spacer

            # Recurse into children
            children = segment.get("children", [])
            if children:
                write_segments(children, depth + 1)

    write_segments(segments)

    # Save to temp file
    report_path = os.path.join(tempfile.gettempdir(), f"codesense_report_{filename}.docx")
    doc.save(report_path)
    return report_path